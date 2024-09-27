import docx.shared
from flask import (render_template, request, 
                   redirect, url_for, flash, session)
from werkzeug.security import generate_password_hash, check_password_hash
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from google.cloud import vision
from dotenv import load_dotenv
import os
import re


from __init__ import app, db
from module import send_word_file, initialize_sessions


load_dotenv(os.path.join(os.path.dirname(__file__), 'ignore', '.env'))


class Word(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(50), nullable=False, unique=True)
    meaning = db.Column(db.Text, nullable=False)


#トップページ
@app.route("/", methods=["GET"])
def home():
    initialize_sessions()
    return render_template("home.html")


#ログインページ
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        correct_username = os.getenv('LOGIN_USERNAME')
        correct_password = os.getenv('LOGIN_PASSWORD')
        username = request.form["username"]
        password = request.form["password"]

        if username == correct_username and password == correct_password:
            session["user_id"] = 1
            flash("ログインに成功しました")
            return redirect(url_for("create_dict"))
        else:
            flash("ユーザー名またはパスワードが間違っています。")
    
    return render_template("login.html")


#ログアウトページ
@app.route("/logout")
def logout():
    session.pop("user_id", None)
    flash("ログアウトしました")
    return redirect(url_for("login"))

#メインページ(用語集に入れる単語の追加とWORD出力)
@app.route("/create_dict", methods=["GET", "POST"])
def create_dict():
    if "user_id" not in session:
        flash("ログインが必要です。")
        return redirect(url_for("login"))
    if request.method == "POST":
        selected_sessions = request.form.getlist("sessions")
        
        for selected_session in selected_sessions: 
            if selected_session == "words":
                session["words"] = session["new_words"] = session["heading"] = []
            elif selected_session == "change_words":
                session["change_words"]  = []
            elif selected_session == "image_texts":
                session["image_texts"] = []
            elif selected_session == "existing_words":
                session["existing_words"] = [word.name for word in Word.query.filter().all()]
    
    words = session["words"]
    return render_template(
        "create_dict.html",
        words=words,
        texts=session["image_texts"]
    )



#DBに無い単語をリストにしてinput_meaningsに遷移する。
@app.route("/add_word", methods=["POST"])
def add_word():


    word_names = request.form.get("word_names")
    word_list = [word.strip() for word in word_names.split("\n") if word.strip()]
    session["words"] = word_list
    for word in word_list:

        if word.startswith("!") or word.startswith("！"):
               session["heading"].append(word) 

        elif (word not in session["existing_words"]) and (word not in session["new_words"]):
            session["new_words"].append(word)
    
    session.modified = True
        
    if session["new_words"]:
        return redirect(url_for("input_meanings"))
    
    else:
        flash("全ての単語がデータベースに存在します。")
        return redirect(url_for("create_dict"))

#DBに無い単語の意味を追加する。
@app.route("/input_meanings", methods=["GET","POST"])
def input_meanings():
    if "user_id" not in session:
        flash("ログインが必要です。")
        return redirect(url_for("login"))
    if request.method == "POST":
        meanings = {}
        new_words = []
        for word in session["new_words"]:
            meaning = request.form.get(f"meaning_{word}")  
            meanings[word] = meaning 
            print(meanings.items())
        for word, meaning in meanings.items():
            if meaning and meaning.strip():
                new_word = Word(name=word, meaning=meaning)
                new_words.append(new_word)
                session["existing_words"].append(word)
        
        session.modified  = True
        db.session.add_all(new_words)
        db.session.commit()
        session.pop("new_words", None)
        flash("意味をデータベースに保存しました。")
        return redirect(url_for("create_dict"))
    return render_template("input_meanings.html", words=session["new_words"])

#wordに出力する
@app.route("/export_word", methods=["POST"])
def export_word():
    word_title = request.form.get("word_title")
    export_treat = request.form.get('export_treat')

    if export_treat == "make_dict":
        words = session["words"]

        words_for_query = [word for word in words if word not in session["heading"]]
        words_from_db = Word.query.filter(Word.name.in_(words_for_query)).all()

        final_export_list = []

        for word in words:
            if word in session["heading"]:
                final_export_list.append(word[1:])

            else:
                db_word = next((w for w in words_from_db if w.name == word), None)
                if db_word:
                    final_export_list.append(db_word)


        doc = docx.Document()
        title_para = doc.add_paragraph(f"{word_title}")
        title_run = title_para.runs[0]
        title_run.font.size = docx.shared.Pt(30)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()
    
        for word in final_export_list:
            
            paragraph = doc.add_paragraph()
            if isinstance(word, Word):
                run_word = paragraph.add_run(f"{word.name}：")
                run_word.font.bold = True

                run_word.font.size = docx.shared.Pt(16)

                # 意味を追加
                run_meaning = paragraph.add_run(f"{word.meaning}")
                run_meaning.font.size = docx.shared.Pt(16)
        
            else:
                run_word = paragraph.add_run(word)
                run_word.font.bold = True
                run_word.font.size = docx.shared.Pt(26)
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph()
    #全てのデータを出力する。
    else:
        all_data = [word for word in Word.query.filter().all()]
        doc = docx.Document()
        title_para = doc.add_paragraph(f"{word_title}")
        title_run = title_para.runs[0]
        title_run.font.size = docx.shared.Pt(20)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()
        
        for word in all_data:
            
            paragraph = doc.add_paragraph()
            if isinstance(word, Word):
                run_word = paragraph.add_run(f"{word.name}：")
                run_word.font.size = docx.shared.Pt(8)

                # 意味を追加
                run_meaning = paragraph.add_run(f"{word.meaning}")
                run_meaning.font.size = docx.shared.Pt(8)
            
  
            doc.add_paragraph()

    # 一時ファイルを作成
    send_word_file(doc)




# 2,WORDの用語をDBに登録する。
@app.route("/word_to_data", methods=["GET", "POST"])
def word_to_data():
    if "user_id" not in session:
        flash("ログインが必要です。")
        return redirect(url_for("login"))
    if request.method == "POST":
        file = request.files["file"]
        word_treat = request.form.get('word_treat')

        if file:
            #ワードの用語をdbに登録する。
            if word_treat == "word_to_data":
                document = docx.Document(file)
                lis = []
                new_words = []
                i = 1
                for paragraph in document.paragraphs:
                    if any(char in paragraph.text for char in ":;：；"):
                        text = "".join(paragraph.text.split())
                        lis.append(text)
                session["existing_words"] = [word.name for word in Word.query.filter().all()]
                for word in lis:
                    word = re.sub(r'^\d+\s*', '', word)
                    word_name, word_meaning = re.split("[:;：；]", word, maxsplit=1)
                    if word_name in session["existing_words"]:
                        continue
                    else:
                        new_word = Word(name=word_name, meaning=word_meaning) 
                        new_words.append(new_word)
                        session["existing_words"].append(word_name)
                        i += 1
                session.modified  = True
                
                if new_words:
                    db.session.add_all(new_words)
                    db.session.commit()
                flash(f"{i}個の単語をデータベースに保存しました。")
                return redirect(url_for("word_to_data"))
            
            #既にできているword用語集を再整形する。
            if word_treat == "word_to_word":
                document = docx.Document(file)
                word_list = []
                for paragraph in document.paragraphs:
                    if any(char in paragraph.text for char in ":;：；"):
                        text = "".join(paragraph.text.split())
                        word = re.sub(r'^\d+\s*', '', text)
                        word_name, word_meaning = re.split("[:;：；]", word, maxsplit=1)
                        word_list.append({"name":word_name, "meaning":word_meaning})
                    else:
                        word_list.append({"heading":paragraph.text})


                doc = docx.Document()
                title_para = doc.add_paragraph(f"{word_list[0]['heading']}")
                title_run = title_para.runs[0]
                title_run.font.size = docx.shared.Pt(30)
                title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                doc.add_paragraph()
                for word_data in word_list:
                    if  word_data.get("name"):
                        paragraph = doc.add_paragraph()

                        run_word = paragraph.add_run(f"{word_data['name']}：")
                        run_word.font.bold = True

                        run_word.font.size = docx.shared.Pt(16)

                        # 意味を追加
                        run_meaning = paragraph.add_run(f"{word_data['meaning']}")
                        run_meaning.font.size = docx.shared.Pt(16)
                        doc.add_paragraph()
                    
                    elif word_data.get("heading"):
                        if word_data == word_list[0]:
                            continue
                        paragraph = doc.add_paragraph()
                        
                        run_heading = paragraph.add_run(word_data["heading"])
                        run_heading.font.bold = True

                        run_heading.font.size = docx.shared.Pt(26)
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                        doc.add_paragraph()
                send_word_file(doc)


    elif request.method == "GET":
        return render_template("word_to_data.html")
    

#DBの意味の変更ページの表示
@app.route("/change_meanings", methods=["GET"])
def change_meanings():
    if "user_id" not in session:
        flash("ログインが必要です。")
        return redirect(url_for("login"))
    return render_template("change_meanings.html")

#送られたフォームから単語がDBに入っているか確認して仕分ける。
@app.route("/change_words", methods=["POST"])
def change_words():
    word_names = request.form.get("change_names")
    word_list = [word.strip() for word in word_names.split("\n") if word.strip()]
    session["new_words"] = []
    session["change_words"] = []
    for word in word_list:
        existing_word = Word.query.filter_by(name=word).first()
        if (existing_word is None) and (word not in session["new_words"]):
            session["new_words"].append(word)
            
        elif word not in session["change_words"] :
            session["change_words"].append({"name":existing_word.name, "meaning":existing_word.meaning})

    session.modified = True
    if session["change_words"] or session["new_words"]:
        return redirect(url_for("input_change_meanings"))

    else:
        return redirect(url_for("change_meanings"))


#意味を変更するフォーム
@app.route("/input_change_meanings", methods=["GET","POST"])
def input_change_meanings():
    if "user_id" not in session:
        flash("ログインが必要です。")
        return redirect(url_for("login"))
    if request.method == "POST":
        #新しい単語を追加
        if session.get("new_words", []):
            meanings = {}
            new_words = []
            for word in session["new_words"]:
                meaning = request.form.get(f"meaning_new_{word}")  
                meanings[word] = meaning 
            for word, meaning in meanings.items():
                if meaning and meaning.strip():
                    new_word = Word(name=word, meaning=meaning)
                    new_words.append(new_word)
            
            db.session.add_all(new_words)
            session.pop("new_words", None)

        #意味を変更する。
        if session.get("change_words", []):
            for word in session["change_words"]:
                name = word["name"]
                new_meaning = request.form.get(f"meaning_change_{word['name']}")
                existing_word = Word.query.filter_by(name=name).first()
                existing_word.meaning = new_meaning
            session.pop("change_words", None)
         
        
        db.session.commit()
        flash(f"単語の意味を更新しました。")

        return redirect(url_for("change_meanings"))
    


    return render_template(
    "input_change_meanings.html",
    change_words=session.get("change_words", []),
    new_words=session.get("new_words", [])
    )



#画像の文章を抽出して表示する。
@app.route("/detect_text", methods=["GET","POST"])
def detect_text():
    if request.method == "POST":
        file = request.files["file"]
        """Detects text in the file."""
        os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = os.path.join(os.path.dirname(__file__), 'ignore', 'vison-app-434901-2321b498d0f2.json')

        client = vision.ImageAnnotatorClient()

        content = file.read()

        image = vision.Image(content=content)

        response = client.text_detection(image=image)
        texts = response.text_annotations
        if "image_texts" not in session:
            session["image_texts"] = []
        session["image_texts"] = []
        if  texts:
            session["image_texts"] = texts[0].description
        
        
        return redirect(url_for("create_dict"))
    
    else:
        return render_template("detect_text.html")
